from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path(r"D:\AIwork\workbuddy\Personal_portfolio\SANJIN-Resume.docx")


def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, 13, True, (181, 82, 45))
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "D66A3B")
    border.append(bottom)
    p._p.get_or_add_pPr().append(border)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.18
    set_font(p.add_run(text))


def project(doc, title, stack, lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(title), 11, True, (35, 44, 55))
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("技术 / 工具：")
    set_font(r, 9.5, True, (99, 112, 128))
    set_font(p.add_run(stack), 9.5, False, (72, 84, 98))
    for line in lines:
        bullet(doc, line)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.45)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run("SANJIN"), 22, True, (35, 44, 55))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_font(p.add_run("应届毕业生 | 数据分析与应用开发方向"), 10.5, False, (99, 112, 128))

    info = doc.add_table(rows=1, cols=1)
    info.autofit = False
    cell = info.cell(0, 0)
    set_cell_shading(cell, "F8EEE8")
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("作品与公开代码：github.com/cuber-sanjin"), 10, True, (160, 68, 34))

    heading(doc, "个人概述")
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.22
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run("应届毕业生，具备 Python 数据处理、Java Web 基础、HTML/CSS/JavaScript 页面开发与 MySQL 基础使用经验。能够在 AI 工具辅助下完成需求拆解、代码理解、运行验证、问题定位和文档整理，重视对生成代码的复核与测试。"))

    heading(doc, "专业技能")
    skills = [
        ("数据分析与可视化", "Python、Pandas、NumPy、jieba、Matplotlib、ECharts；了解情感分类、LDA 主题分析、关键词与词云生成。"),
        ("Web 与数据库基础", "了解 Flask 路由、模板与接口交互；掌握 Java 基础、Servlet/JDBC 分层思路；能够使用 MySQL 完成表结构设计与常用 SQL。"),
        ("工程工具", "使用 Git/GitHub 管理代码；使用 PyCharm、VS Code、Android Studio 进行开发和调试；了解 Apifox/Postman 接口验证。"),
    ]
    for name, text in skills:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        set_font(p.add_run(f"{name}："), 10.5, True, (35, 44, 55))
        set_font(p.add_run(text))

    heading(doc, "项目经历")
    project(doc, "基于微博数据的公共事件情感倾向分析系统（毕业设计）", "Python、Flask、MySQL、BERT、LDA、jieba、ECharts", [
        "围绕公共事件评论数据完成“数据清洗 - 中文分词 - 情感分类 - 主题分析 - 可视化展示”的分析流程。",
        "参与 Flask 接口、数据处理脚本、MySQL 数据表、前端页面与图表展示的联调；使用 ECharts 呈现评论量、情感占比与趋势。",
        "基于 BERT 思路进行情感倾向分类，并结合 LDA、关键词和词云辅助观察讨论主题；本地演示使用脱敏样例数据。",
    ])
    project(doc, "图书借阅管理系统（Java Web 课程设计）", "Java、Servlet、JDBC、MySQL、Maven、Tomcat、Bootstrap", [
        "梳理管理员和读者两类角色的图书管理、读者管理、借还记录查询、图书查询和个人信息维护等页面流程。",
        "阅读并验证 Servlet - Service - DAO - POJO 分层结构，了解 JDBC、c3p0 连接池与 MySQL CRUD 的协作方式。",
        "项目保留原始课程设计源码；需连接本地数据库运行，不对外提供账号、数据库或源码下载。",
    ])
    project(doc, "痛风患者每日食谱（微信小程序）", "WXML/WXSS/JavaScript、CloudBase、Canvas、多模态 AI", [
        "完成食物嘌呤查询、尿酸记录、饮水与运动建议等页面功能；以食物数据分类和风险分级支持日常查询。",
        "尝试接入图片识别能力，并设计识别不可用时转为手动输入和文本匹配的降级路径。",
    ])
    project(doc, "移动记账应用原型", "HTML/CSS/JavaScript、PWA、Capacitor、本地存储", [
        "实现收入支出记录、分类管理、金额统计和历史记录查看等基础功能，了解表单、列表和本地数据持久化。",
        "使用 AI 辅助完成交互逻辑与异常处理的验证，并根据运行结果调整页面与功能。",
    ])
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("SANJIN | 个人项目简历"), 8.5, False, (130, 140, 150))
    doc.core_properties.author = "SANJIN"
    doc.core_properties.last_modified_by = "SANJIN"
    doc.core_properties.title = "SANJIN Resume"
    doc.save(OUT)


if __name__ == "__main__":
    main()
